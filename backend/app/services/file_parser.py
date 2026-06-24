"""
Extracts plain text from uploaded files.

Supported: PDF (PyMuPDF), DOCX (python-docx), TXT / plain text.
Returned text is stripped of excess whitespace but otherwise unmodified.
"""

from __future__ import annotations

import io
import logging
import zipfile

logger = logging.getLogger(__name__)

# Maximum characters to pass to the LLM (≈30 k tokens at ~4 chars/token)
MAX_CHARS = 120_000


def extract_text(content: bytes, filename: str) -> str:
    """Dispatch to the appropriate parser based on file extension."""
    name = filename.lower()
    if name.endswith(".pdf"):
        return _pdf(content)
    if name.endswith(".docx"):
        return _docx(content)
    # Fallback: treat as plain text (UTF-8, fallback to latin-1)
    return _plain(content)


# ── PDF ───────────────────────────────────────────────────────────────────────

def _pdf(content: bytes) -> str:
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(stream=content, filetype="pdf")
        pages: list[str] = []
        for page in doc:
            pages.append(page.get_text("text"))
        text = "\n".join(pages)
        logger.info("PDF: extracted %d chars from %d pages", len(text), len(doc))
        return _clean(text)
    except ImportError:
        raise RuntimeError("PyMuPDF is not installed (pip install PyMuPDF)")
    except Exception as exc:
        raise RuntimeError(f"PDF extraction failed: {exc}") from exc


# ── DOCX ──────────────────────────────────────────────────────────────────────

_MAX_DOCX_DECOMPRESSED = 50 * 1024 * 1024  # 50 MB


def _docx(content: bytes) -> str:
    # Guard against zip bombs: check total decompressed size before python-docx
    # opens the archive, since python-docx doesn't impose a size limit itself.
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as zf:
            total = sum(info.file_size for info in zf.infolist())
            if total > _MAX_DOCX_DECOMPRESSED:
                raise RuntimeError(
                    f"DOCX decompressed size {total} bytes exceeds {_MAX_DOCX_DECOMPRESSED} limit"
                )
    except zipfile.BadZipFile as exc:
        raise RuntimeError(f"DOCX is not a valid ZIP archive: {exc}") from exc

    try:
        from docx import Document

        doc = Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also pull text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        text = "\n".join(paragraphs)
        logger.info("DOCX: extracted %d chars", len(text))
        return _clean(text)
    except ImportError:
        raise RuntimeError("python-docx is not installed (pip install python-docx)")
    except Exception as exc:
        raise RuntimeError(f"DOCX extraction failed: {exc}") from exc


# ── Plain text ────────────────────────────────────────────────────────────────

def _plain(content: bytes) -> str:
    for enc in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            return _clean(content.decode(enc))
        except (UnicodeDecodeError, ValueError):
            continue
    return _clean(content.decode("utf-8", errors="replace"))


# ── Helpers ───────────────────────────────────────────────────────────────────

def _clean(text: str) -> str:
    import re
    # Collapse runs of blank lines to a single blank line
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Collapse runs of spaces/tabs to a single space
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()[:MAX_CHARS]
