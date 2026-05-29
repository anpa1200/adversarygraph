"""Unit tests for the file parser service.  No external dependencies."""

import io
import pytest


# ── Plain text ─────────────────────────────────────────────────────────────────

def test_plain_text_utf8():
    from app.services.file_parser import extract_text
    content = "APT29 used spearphishing attachments (T1566.001) to gain access.".encode("utf-8")
    result = extract_text(content, "report.txt")
    assert "APT29" in result
    assert "T1566.001" in result


def test_plain_text_latin1():
    from app.services.file_parser import extract_text
    content = "Adversary used credential dumping (T1003).".encode("latin-1")
    result = extract_text(content, "notes.txt")
    assert "T1003" in result


def test_plain_text_excess_whitespace_collapsed():
    from app.services.file_parser import extract_text
    content = "Line one.\n\n\n\n\nLine two.".encode("utf-8")
    result = extract_text(content, "file.txt")
    assert "\n\n\n" not in result


def test_plain_text_truncated_at_max():
    from app.services.file_parser import extract_text, MAX_CHARS
    content = ("A" * (MAX_CHARS + 500)).encode("utf-8")
    result = extract_text(content, "big.txt")
    assert len(result) <= MAX_CHARS


# ── DOCX ──────────────────────────────────────────────────────────────────────

def test_docx_extraction():
    """Build a real DOCX in memory and verify extraction."""
    pytest.importorskip("docx")
    from docx import Document
    from app.services.file_parser import extract_text

    doc = Document()
    doc.add_paragraph("This is a threat intelligence report.")
    doc.add_paragraph("Technique: T1059 — Command and Scripting Interpreter.")

    buf = io.BytesIO()
    doc.save(buf)
    content = buf.getvalue()

    result = extract_text(content, "report.docx")
    assert "T1059" in result
    assert "threat intelligence" in result.lower()


def test_docx_table_extraction():
    pytest.importorskip("docx")
    from docx import Document
    from app.services.file_parser import extract_text

    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Technique"
    table.cell(0, 1).text = "Tactic"
    table.cell(1, 0).text = "T1566"
    table.cell(1, 1).text = "initial-access"

    buf = io.BytesIO()
    doc.save(buf)
    result = extract_text(buf.getvalue(), "table.docx")
    assert "T1566" in result


# ── PDF ───────────────────────────────────────────────────────────────────────

def test_pdf_extraction():
    """Build a minimal PDF with PyMuPDF and verify extraction."""
    fitz = pytest.importorskip("fitz")
    from app.services.file_parser import extract_text

    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 100), "Lateral movement via T1021.001 (RDP).")
    buf = io.BytesIO()
    doc.save(buf)

    result = extract_text(buf.getvalue(), "sample.pdf")
    assert "T1021" in result


# ── Edge cases ─────────────────────────────────────────────────────────────────

def test_empty_content():
    from app.services.file_parser import extract_text
    result = extract_text(b"", "empty.txt")
    assert result == ""


def test_unknown_extension_falls_back_to_text():
    from app.services.file_parser import extract_text
    content = "Some text content.".encode("utf-8")
    result = extract_text(content, "file.log")
    assert "Some text content" in result
